'''
Autor: Luiz Fernando Antonelli Galati
Julho/2023
'''

'''
Este código lê vários arquivos docx (Word), cada um deles contendo uma tabela intitulada "Objetivos do 
programa de graduação FGV Direito SP", e transfere o conteúdo de cada uma dessas tabelas para um único
arquivo xlsx (Excel) intitulado "Dados por disciplina...". Mais simplificadamente, esse programa transfere
várias tabelas de diferentes arquivos Word para um único arquivo Excel.
'''



import docx
import os
import xlsxwriter
import time

def escreveQuesitos (novaPlanilha):
    novaPlanilha.write (1, 0, "Objetivos do programa de graduação FGV Direito SP")
    novaPlanilha.write (2, 0, "Domínio de conceitos, estruturas e racionalidades fundamentais do Direito")
    novaPlanilha.write (3, 0, "Conhecimento de áreas contíguas ao Direito")
    novaPlanilha.write (4, 0, "Aplicação prática de conceitos e estruturas do Direito")
    novaPlanilha.write (5, 0, "Pesquisa Jurídica")
    novaPlanilha.write (6, 0, "Comunicação")
    novaPlanilha.write (7, 0, "Colaboração e trabalho em rede")
    novaPlanilha.write (8, 0, "Ética")
    novaPlanilha.write (9, 0, "Empreendedorismo")
    novaPlanilha.write (10, 0, "Cosmopolitanismo")
    novaPlanilha.write (11, 0, "Outros objetivos da disciplina")


def verificaTabela (table):
    nLinhas = len (table.rows)
    nColunas = len (table.columns)
    if (nLinhas != 11 or nColunas != 3):
        return 0    

    textA = table.cell(0, 0).text  
    if (textA != "Objetivos do programa de graduação FGV Direito SP"):
        return 0

    return 1


def transfereTabela (table, novaPlanilha, i, k):
    l = 0
    while (l < len (table.rows)):
        textA = table.cell(l, 1).text
        textB = table.cell(l, 2).text
        
        if (l == 0): # linha é cabeçalho. apenas copia o conteúdo da terceira coluna.
            novaPlanilha.write (l + 1, k, textA)
            novaPlanilha.write (l + 1, k + 1, textB)
        elif (l >= 1 and l <= 9):
            novaPlanilha.write (l + 1, k, textA)
            if (textB == "○ ○ ○" or textB == "0"):                
                novaPlanilha.write (l + 1, k + 1, 0)
            elif (textB == "● ○ ○" or textB == "1"):
                novaPlanilha.write (l + 1, k + 1, 1)            
            elif (textB == "● ● ○" or textB == "○●●" or textB == "2"):
                novaPlanilha.write (l + 1, k + 1, 2)
            elif (textB == "● ● ●" or textB == "●●●" or textB == "3"):
                novaPlanilha.write (l + 1, k + 1, 3)           
            else:
                novaPlanilha.write (l + 1, k + 1, "Erro")
        else:
            if (textB != "Outros objetivos da disciplina: ---" and textB != "Outros objetivos da disciplina: "):
                novaPlanilha.write (l + 1, k, textB)
        l = l + 1


def main ():
    diretorioAtual = os.getcwd ()

    diretorioPastasSemestres = diretorioAtual + "/Obrigatórias"
    if (os.path.isdir (diretorioPastasSemestres)):
        listaSemestres = os.listdir (diretorioPastasSemestres)
        novoArquivo = xlsxwriter.Workbook ("Dados por disciplina (obrigatórias) - Revisar2.xlsx")

        j = 0
        while (j < len (listaSemestres)):
            novaPlanilha = novoArquivo.add_worksheet (listaSemestres[j])
        
            listaDisciplinas = os.listdir (diretorioPastasSemestres + "/" + listaSemestres[j])
            i = 0           # número do arquivo da pasta de um semestre específico (= é o primeiro arquivo (arquivo 0)? é o segundo (arquivo 1?))
            k = 1           # coluna a ser utilizada para escrever no arquivo xlsx  
            while (i < len (listaDisciplinas)):
                novaPlanilha.write (0, k, listaDisciplinas[i])
                doc = docx.Document (diretorioPastasSemestres + "/" + listaSemestres[j] + "/" + listaDisciplinas[i])
                escreveQuesitos (novaPlanilha)

                l = 0
                while (l < len (doc.tables)):
                    table = doc.tables[l]
                    if (verificaTabela (table) == 1): # testa a tabela para ver se ela está dentro dos padrões               
                        transfereTabela (table, novaPlanilha, i, k)
                        break
                    l = l + 1
                if (l == len (doc.tables)):                  # programa não encontrou a tabela
                    novaPlanilha.write (1, k, "Objetivos da disciplina")
                    novaPlanilha.write (1, k + 1, "Grau de contribuição")
                    novaPlanilha.write (2, k, "erro!")
                    novaPlanilha.write (2, k + 1, "erro!")

                i = i + 1
                k = k + 2
        
            j = j + 1

        novoArquivo.close ()
    else:
        print ("A pasta das disciplinas obrigatórias não está disponível!")
        time.sleep (1.5)       

    diretorioPastasSemestres = diretorioAtual + "/Eletivas e clínicas"
    if (os.path.isdir (diretorioPastasSemestres)):
        listaSemestres = os.listdir (diretorioPastasSemestres)
        novoArquivo = xlsxwriter.Workbook ("Dados por disciplina (eletivas e clínicas) - Revisar2.xlsx")

        j = 0
        while (j < len (listaSemestres)):
            novaPlanilha = novoArquivo.add_worksheet (listaSemestres[j])
        
            listaDisciplinas = os.listdir (diretorioPastasSemestres + "/" + listaSemestres[j])
            i = 0           # número do arquivo da pasta de um semestre específico (= é o primeiro arquivo (arquivo 0)? é o segundo (arquivo 1?))
            k = 1           # coluna a ser utilizada para escrever no arquivo xlsx  
            while (i < len (listaDisciplinas)):
                novaPlanilha.write (0, k, listaDisciplinas[i])
                doc = docx.Document (diretorioPastasSemestres + "/" + listaSemestres[j] + "/" + listaDisciplinas[i])
                escreveQuesitos (novaPlanilha)

                l = 0
                while (l < len (doc.tables)):
                    table = doc.tables[l]
                    if (verificaTabela (table) == 1): # testa a tabela para ver se ela está dentro dos padrões               
                        transfereTabela (table, novaPlanilha, i, k)
                        break
                    l = l + 1
                if (l == len (doc.tables)):                  # programa não encontrou a tabela
                    novaPlanilha.write (1, k, "Objetivos da disciplina")
                    novaPlanilha.write (1, k + 1, "Grau de contribuição")
                    novaPlanilha.write (2, k, "erro!")
                    novaPlanilha.write (2, k + 1, "erro!")

                i = i + 1
                k = k + 2
        
            j = j + 1   
        
        novoArquivo.close ()
    else:
        print ("A pasta das disciplinas eletivas e das clínicas não está disponível!")
        time.sleep (1.5)       

main ()
